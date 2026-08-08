"""Work 模式业务路由（Task 13 / 14 / 15 / 16）。

集中提供 Work 模式下的工作对象抽取与入图、行业风口推荐、工作报告生成与导出、
用户提问回答接口，挂载在 ``/api`` 前缀下：

- ``POST /api/graphs/{graph_id}/work/extract``
  从用户输入文本抽取候选工作对象（不入图，返回带关系信息的候选列表）。
- ``POST /api/graphs/{graph_id}/work/confirm``
  批量确认入图（归一去重，按 relation 建立边）。
- ``POST /api/graphs/{graph_id}/work/trends``
  基于当前 work 图谱分析并生成行业风口推荐。
- ``POST /api/graphs/{graph_id}/work/trends/{index}/add-to-graph``
  把指定风口转为图谱节点（复用工作线索类型）。
- ``POST /api/graphs/{graph_id}/work/report``
  生成结构化工作报告（Markdown + sections）。
- ``POST /api/graphs/{graph_id}/work/report/export-docx``
  把 Markdown 转为 .docx 文件流返回。
- ``POST /api/graphs/{graph_id}/work/ask``
  基于图谱上下文回答用户提问，标注来源与置信度。

设计要点：

1. **不修改 graph_store / graph_agent / node_types**：仅组合调用既有方法，
   与 extensions.py / extraction.py / quiz.py 路由风格一致。
2. **关系映射到 EDGE_RELATIONS 语义**：confirm 接口在创建节点后按
   ``relation`` 枚举（belongs_to/involves/committed_to/depends_on/waiting_for/
   influences/source_of/alternative_to/related）建立边；``to_title`` 可能指向
   本次新创建节点或已存在节点，二者都通过标题查表回填 id。
3. **降级透明传递**：所有 Agent 方法在 LLM 不可用时返回降级结构，本层原样
   透传，前端据此显示「AI 服务暂不可用」提示但不阻断流程。
4. **docx 导出走临时文件**：用 ``python-docx`` 生成 .docx 后通过
   ``StreamingResponse`` 返回，文件名按报告周期 + 时间戳生成，避免冲突。
5. **流式版本简化**：当前版本非流式（一次性返回结果）。流式入口已在
   ``graph_agent.*_stream`` 中实现并通过 ``ws_notify`` 推送，本路由层后续可
   通过 ``GET /work/.../stream`` 形式追加（依赖 WS session 注册）。
"""

from __future__ import annotations

import io
import logging
import re
import uuid
from typing import Any, Literal

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.models.node_types import (
    EDGE_RELATED,
    GRAPH_TYPE_WORK,
    NODE_SOURCE_AGENT,
    WORK_OBJECT_THREAD,
)
from app.services.graph_agent import GraphAgent, _titles_similar, get_graph_agent
from app.services.graph_store import GraphStore, graph_store

logger = logging.getLogger(__name__)

router = APIRouter()


# ============================================================================
# 进程内缓存：风口推荐结果（add-to-graph 按 index 取回）
# ============================================================================

#: ``graph_id -> [trend, ...]``，仅缓存最近一次生成的风口列表。
#:
#: 重启丢失可接受；前端可随时调 trends 接口重新生成刷新缓存。
_trends_cache: dict[str, list[dict[str, Any]]] = {}


# ============================================================================
# 依赖注入
# ============================================================================


def get_graph_store_dep() -> GraphStore:
    """依赖注入：返回全局 GraphStore 单例。"""
    return graph_store


def get_agent() -> GraphAgent:
    """依赖注入：返回全局 GraphAgent 单例。"""
    return get_graph_agent()


def _not_found(detail: str) -> HTTPException:
    return HTTPException(status_code=404, detail=detail)


def _bad_request(detail: str) -> HTTPException:
    return HTTPException(status_code=400, detail=detail)


async def _ensure_work_graph(graph_id: str, store: GraphStore) -> dict[str, Any]:
    """校验图谱存在且为 work 模式，返回图谱 dict。"""
    graph = await store.get_graph(graph_id)
    if graph is None:
        raise _not_found(f"图谱不存在: {graph_id}")
    if graph.get("type") != GRAPH_TYPE_WORK:
        raise _bad_request(f"该接口仅支持 work 图谱: {graph_id}")
    return graph


# ============================================================================
# 请求 / 响应模型
# ============================================================================


class WorkExtractRequest(BaseModel):
    """从文本抽取工作对象请求。"""

    text: str = Field(..., min_length=1, description="用户输入的工作信息文本")


class WorkRelationItem(BaseModel):
    """工作对象间的关系（confirm 时携带）。"""

    to_title: str = Field(..., description="关系目标对象的标题")
    relation: str = Field(
        "related",
        description=(
            "关系语义：related/belongs_to/involves/committed_to/depends_on/"
            "waiting_for/influences/source_of/alternative_to"
        ),
    )


class WorkObjectItem(BaseModel):
    """单个候选工作对象（待确认或已确认入图）。"""

    title: str = Field(..., min_length=1, max_length=255)
    summary: str = Field("", description="一句话概括")
    type: str = Field(WORK_OBJECT_THREAD, description="工作对象子类型")
    relations: list[WorkRelationItem] = Field(default_factory=list)


class WorkConfirmRequest(BaseModel):
    """批量确认入图请求。"""

    objects: list[WorkObjectItem] = Field(..., min_length=1)


class WorkConfirmResponse(BaseModel):
    """批量确认入图响应。"""

    created: list[dict[str, Any]] = Field(default_factory=list)
    skipped: list[dict[str, Any]] = Field(default_factory=list)
    edges_created: int = 0
    created_count: int = 0
    skipped_count: int = 0


class TrendsResponse(BaseModel):
    """风口推荐响应。"""

    trends: list[dict[str, Any]] = Field(default_factory=list)
    degraded: bool = False
    cached: bool = False


class TrendAddResponse(BaseModel):
    """风口加入图谱响应。"""

    node: dict[str, Any]
    trend_title: str


class ReportRequest(BaseModel):
    """工作报告生成请求。"""

    period: Literal["weekly", "monthly"] = Field(
        "weekly", description="报告周期：weekly 周报 / monthly 月报"
    )


class ReportResponse(BaseModel):
    """工作报告响应。"""

    markdown: str
    sections: dict[str, list[str]]
    period: str
    degraded: bool = False
    degrade_reason: str = ""


class AskRequest(BaseModel):
    """用户提问请求。"""

    question: str = Field(..., min_length=1, description="用户提问")


class AskResponse(BaseModel):
    """用户提问响应。"""

    answer: str
    sources: list[dict[str, Any]] = Field(default_factory=list)
    confidence: float
    degraded: bool = False
    degrade_reason: str = ""


# ============================================================================
# Task 13：工作对象抽取与入图
# ============================================================================


@router.post(
    "/graphs/{graph_id}/work/extract",
    response_model=dict,
)
async def extract_work_objects(
    graph_id: str,
    body: WorkExtractRequest,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> dict[str, Any]:
    """从用户输入文本抽取候选工作对象（不入图）。

    调用 ``graph_agent.extract_work_objects(text, graph_id)``，返回候选对象列表，
    每项含 ``title / summary / type / relations``（relations 含 ``to_title`` 与
    ``relation``，由前端展示供用户确认。

    LLM 不可用时返回空列表与 ``degraded=True`` 标记，前端据此显示降级提示。
    """
    await _ensure_work_graph(graph_id, store)

    try:
        objects = await agent.extract_work_objects(body.text, graph_id)
    except Exception as exc:  # noqa: BLE001
        logger.warning("work extract: agent 异常: %s", exc)
        objects = []

    degraded = len(objects) == 0
    return {
        "graph_id": graph_id,
        "objects": objects,
        "degraded": degraded,
    }


@router.post(
    "/graphs/{graph_id}/work/confirm",
    response_model=WorkConfirmResponse,
)
async def confirm_work_objects(
    graph_id: str,
    body: WorkConfirmRequest,
    store: GraphStore = Depends(get_graph_store_dep),
) -> WorkConfirmResponse:
    """批量确认工作对象入图（归一去重 + 建立关系边）。

    流程：
    1. 校验图谱存在且为 work 模式。
    2. 加载现有节点，对每个候选对象标题做相似度判断：
       - 与现有节点相似 → 跳过，记录 ``existing_node_id``。
       - 否则调 ``create_node`` 落库，``source=agent``。
    3. 收集本次新建 + 已存在节点，构建 ``title -> node_id`` 索引。
    4. 对每个对象的 ``relations``，按 ``to_title`` 查索引得到目标节点 id，
       调 ``create_edge`` 建立边（``relation`` 校验后透传，非法值兜底为 related）。
    5. 返回 ``created`` / ``skipped`` / ``edges_created`` 等统计。
    """
    await _ensure_work_graph(graph_id, store)

    existing_nodes = await store.list_nodes(graph_id)

    # 标题 -> 节点 id 索引（含现有 + 本批新建）
    title_to_id: dict[str, str] = {}
    for n in existing_nodes:
        title_to_id[n.get("title", "")] = n.get("id", "")

    created: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []
    valid_relations = {
        "related", "belongs_to", "involves", "committed_to", "depends_on",
        "waiting_for", "influences", "source_of", "alternative_to",
        "prerequisite", "extends",
    }

    for obj in body.objects:
        title = obj.title.strip()
        if not title:
            continue

        # 归一去重：与现有节点（含本批已建）标题相似则跳过
        dup_id: str | None = None
        for t, nid in title_to_id.items():
            if t and _titles_similar(t, title):
                dup_id = nid
                break
        if dup_id:
            skipped.append({"title": title, "existing_node_id": dup_id})
            continue

        # 归一 relation
        try:
            node = await store.create_node(
                graph_id=graph_id,
                node_type=obj.type,
                title=title,
                summary=obj.summary,
                detail_payload=None,  # 按工作模板自动初始化空值
                is_gray=False,
                source=NODE_SOURCE_AGENT,
                confidence=0.85,
            )
        except ValueError as exc:
            skipped.append({"title": title, "error": str(exc)})
            continue

        created.append(node)
        title_to_id[node["title"]] = node["id"]

    # 建立关系边：遍历所有对象（含已存在节点）的 relations
    edges_created = 0
    for obj in body.objects:
        src_title = obj.title.strip()
        src_id = title_to_id.get(src_title, "")
        if not src_id:
            continue
        for rel in obj.relations:
            to_title = rel.to_title.strip()
            if not to_title:
                continue
            dst_id = title_to_id.get(to_title, "")
            if not dst_id:
                # 目标不在本次入图范围：尝试在现有节点中模糊匹配
                for t, nid in title_to_id.items():
                    if t and _titles_similar(t, to_title):
                        dst_id = nid
                        break
            if not dst_id or dst_id == src_id:
                continue
            relation = rel.relation if rel.relation in valid_relations else EDGE_RELATED
            try:
                await store.create_edge(
                    graph_id=graph_id,
                    src_id=src_id,
                    dst_id=dst_id,
                    relation=relation,
                )
                edges_created += 1
            except ValueError as exc:
                logger.warning(
                    "work confirm: create_edge 失败 src=%s dst=%s rel=%s err=%s",
                    src_id, dst_id, relation, exc,
                )

    return WorkConfirmResponse(
        created=created,
        skipped=skipped,
        edges_created=edges_created,
        created_count=len(created),
        skipped_count=len(skipped),
    )


# ============================================================================
# Task 14：行业风口推荐
# ============================================================================


@router.post(
    "/graphs/{graph_id}/work/trends",
    response_model=TrendsResponse,
)
async def generate_trends(
    graph_id: str,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> TrendsResponse:
    """基于当前 work 图谱生成行业风口推荐。

    调用 ``graph_agent.generate_trends(graph_id)``，返回风口列表
    ``[{title, reason, relevance, suggested_actions}]``。同时把结果缓存到
    进程内 ``_trends_cache``，供后续 ``add-to-graph`` 按 index 取回。
    """
    await _ensure_work_graph(graph_id, store)

    try:
        trends = await agent.generate_trends(graph_id)
    except Exception as exc:  # noqa: BLE001
        logger.warning("work trends: agent 异常: %s", exc)
        trends = []

    # 更新缓存
    _trends_cache[graph_id] = trends

    return TrendsResponse(
        trends=trends,
        degraded=len(trends) == 0,
        cached=False,
    )


@router.post(
    "/graphs/{graph_id}/work/trends/{index}/add-to-graph",
    response_model=TrendAddResponse,
    status_code=201,
)
async def add_trend_to_graph(
    graph_id: str,
    index: int,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> TrendAddResponse:
    """把指定风口转为图谱节点（复用工作线索类型）。

    若缓存中无对应 index（缓存过期或前端未先调 trends），则重新生成风口列表
    后再取回，避免因缓存失效导致加入失败。
    """
    await _ensure_work_graph(graph_id, store)

    trends = _trends_cache.get(graph_id)
    if not trends:
        # 缓存失效：重新生成
        try:
            trends = await agent.generate_trends(graph_id)
        except Exception as exc:  # noqa: BLE001
            logger.warning("add_trend_to_graph: 重新生成趋势异常: %s", exc)
            trends = []
        _trends_cache[graph_id] = trends

    if index < 0 or index >= len(trends):
        raise _not_found(f"风口索引不存在: {index}（当前共 {len(trends)} 个）")

    trend = trends[index]
    title = (trend.get("title") or "").strip()
    if not title:
        raise _bad_request("风口标题为空，无法加入图谱")

    reason = (trend.get("reason") or "").strip()
    actions = trend.get("suggested_actions") or []
    relevance = trend.get("relevance", "medium")

    # 构造 detail_payload：把风口理由与建议行动写入工作线索模板字段
    actions_text = "\n".join(f"- {a}" for a in actions if a) if actions else ""
    detail_payload = {
        "summary": f"行业风口：{title}（相关度：{relevance}）",
        "key_info": f"相关度：{relevance}\n来源：行业风口推荐",
        "related_persons": "",
        "risks": "",
        "extensions": actions_text or "（暂无建议行动）",
    }
    # 把理由也写入 extensions 字段，供详情卡展示
    if reason:
        detail_payload["extensions"] = (
            f"推荐理由：{reason}\n建议行动：\n{actions_text}"
            if actions_text
            else f"推荐理由：{reason}"
        )

    # 归一去重：若已存在同名节点则不重复创建
    existing_nodes = await store.list_nodes(graph_id)
    for n in existing_nodes:
        if _titles_similar(n.get("title", ""), title):
            return TrendAddResponse(node=n, trend_title=title)

    try:
        node = await store.create_node(
            graph_id=graph_id,
            node_type=WORK_OBJECT_THREAD,
            title=title[:255],
            summary=detail_payload["summary"],
            detail_payload=detail_payload,
            is_gray=False,
            source=NODE_SOURCE_AGENT,
            confidence=0.8,
        )
    except ValueError as exc:
        raise _bad_request(f"创建风口节点失败: {exc}") from exc

    return TrendAddResponse(node=node, trend_title=title)


# ============================================================================
# Task 15：工作报告生成与导出
# ============================================================================


@router.post(
    "/graphs/{graph_id}/work/report",
    response_model=ReportResponse,
)
async def generate_report(
    graph_id: str,
    body: ReportRequest,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> ReportResponse:
    """生成结构化工作报告。

    调用 ``graph_agent.generate_report(graph_id, period)``，返回
    ``{markdown, sections, period, degraded}``。LLM 不可用时返回兜底报告。
    """
    await _ensure_work_graph(graph_id, store)

    try:
        result = await agent.generate_report(graph_id, body.period)
    except Exception as exc:  # noqa: BLE001
        logger.warning("work report: agent 异常: %s", exc)
        result = {
            "markdown": f"# 工作报告\n\n（报告生成服务异常：{exc}）\n",
            "sections": {"progress": [], "plan": [], "risks": [], "commitments": []},
            "period": body.period,
            "degraded": True,
            "degrade_reason": str(exc),
        }

    return ReportResponse(
        markdown=result.get("markdown", ""),
        sections=result.get("sections", {}) or {
            "progress": [], "plan": [], "risks": [], "commitments": [],
        },
        period=result.get("period", body.period),
        degraded=bool(result.get("degraded")),
        degrade_reason=result.get("degrade_reason", ""),
    )


@router.post(
    "/graphs/{graph_id}/work/report/export-docx",
)
async def export_report_docx(
    graph_id: str,
    body: ReportRequest,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> StreamingResponse:
    """把工作报告转为 .docx 文件流返回。

    复用步影已依赖的 ``python-docx``（pyproject.toml 已声明 ``python-docx>=1.1.2``）。
    流程：
    1. 调 ``generate_report`` 拿到 Markdown 文本。
    2. 解析 Markdown（标题 ``#`` / ``##`` / 段落 / ``-`` 列表）逐行写入 docx。
    3. 用 ``StreamingResponse`` 返回内存字节流，文件名按周期 + 时间戳生成。
    """
    from docx import Document  # type: ignore[import-not-found]

    await _ensure_work_graph(graph_id, store)

    try:
        result = await agent.generate_report(graph_id, body.period)
    except Exception as exc:  # noqa: BLE001
        logger.warning("export_report_docx: generate_report 异常: %s", exc)
        result = {
            "markdown": f"# 工作报告\n\n（报告生成服务异常：{exc}）\n",
            "period": body.period,
            "degraded": True,
        }

    markdown = result.get("markdown", "") or ""

    doc = Document()
    # 解析 Markdown 写入 docx（支持 # / ## / ### 标题、- 列表、空行分段）
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            # 空行：插入空段落以保留结构
            doc.add_paragraph("")
            continue
        # 标题
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\s*[-*]\s+", line):
            # 无序列表项
            content = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(content, style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            # 有序列表项
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            doc.add_paragraph(content, style="List Number")
        else:
            doc.add_paragraph(line)

    # 写入内存字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    file_id = uuid.uuid4().hex[:8]
    filename = f"work_{body.period}_{file_id}.docx"
    # 中文文件名需 RFC5987 编码，避免 Content-Disposition 乱码
    encoded = filename.encode("utf-8").decode("ascii", errors="ignore") or file_id

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f'attachment; filename="work_{body.period}_{file_id}.docx"; '
                f"filename*=UTF-8''{encoded}"
            ),
        },
    )


# ============================================================================
# Task 16：用户提问回答
# ============================================================================


@router.post(
    "/graphs/{graph_id}/work/ask",
    response_model=AskResponse,
)
async def ask_question(
    graph_id: str,
    body: AskRequest,
    store: GraphStore = Depends(get_graph_store_dep),
    agent: GraphAgent = Depends(get_agent),
) -> AskResponse:
    """基于工作图谱上下文回答用户提问。

    调用 ``graph_agent.answer_question(graph_id, question)``，返回
    ``{answer, sources, confidence, degraded}``。LLM 不可用时返回兜底回答。
    """
    await _ensure_work_graph(graph_id, store)

    try:
        result = await agent.answer_question(graph_id, body.question)
    except Exception as exc:  # noqa: BLE001
        logger.warning("work ask: agent 异常: %s", exc)
        result = {
            "answer": f"（问答服务异常：{exc}）",
            "sources": [],
            "confidence": 0.0,
            "degraded": True,
            "degrade_reason": str(exc),
        }

    return AskResponse(
        answer=result.get("answer", ""),
        sources=result.get("sources", []) or [],
        confidence=float(result.get("confidence", 0.0)),
        degraded=bool(result.get("degraded")),
        degrade_reason=result.get("degrade_reason", ""),
    )
